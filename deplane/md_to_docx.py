from markdown import markdown as md
from lxml import etree
import docx
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


def insert_markdown(document, markdown):
    """
    Insert markdown into docx document, best effort conversion
    """
    html = md(
        markdown,
        extensions=['tables'],
    )
    root = etree.HTML(html)

    for element in root[0]:
        if element.tag == 'h1':
            document.add_heading(element.text, 1)
        elif element.tag == 'h2':
            document.add_heading(element.text, 2)
        elif element.tag == 'h3':
            document.add_heading(element.text, 3)

        elif element.tag == 'p':
            p = document.add_paragraph(element.text.strip())
            finish_paragraph(p, element)

        elif element.tag == 'ul':
            for li in element:
                p = document.add_paragraph(li.text, style='List Bullet')
                finish_paragraph(p, li)

        elif element.tag == 'ol':
            for li in element:
                p = document.add_paragraph(li.text, style='List Number')
                finish_paragraph(p, li)

        elif element.tag == 'table':
            thead = element[0]
            table = document.add_table(rows=0, cols=len(thead[0]))
            for td, cell in zip(thead[0], table.add_row().cells):
                cell.text = td.text

            tbody = element[1]
            for tr in tbody:
                for td, cell in zip(tr, table.add_row().cells):
                    cell.text = td.text

            format_table(table, top_color='d9d9d9')

        else:
            assert 0, element.tag


def format_table(table, widths=None, top_color=None, left_color=None):
    if widths:
        # for Word
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = w
        # for Libreoffice
        for i, w in enumerate(widths):
            table.columns[i].width = w

    if left_color:
        for cell in table.columns[0].cells:
            element = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), left_color))
            cell._tc.get_or_add_tcPr().append(element)

    if top_color:
        for cell in table.rows[0].cells:
            element = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), top_color))
            cell._tc.get_or_add_tcPr().append(element)


def finish_paragraph(p, element):
    """
    add children of element as runs at the end of p, best effort
    """
    for sub in element:
        if sub.tag == 'em':
            p.add_run(sub.text).italic = True
        elif sub.tag == 'b':
            p.add_run(sub.text).bold = True
        elif sub.tag == 'a':
            add_hyperlink(p, sub.attrib['href'], sub.text)
        else:
            p.add_run(sub.text)
        if sub.tail:
            p.add_run(sub.tail)


def add_hyperlink(paragraph, url, text):
    """
    A function that places a hyperlink within a paragraph object.
    from https://github.com/python-openxml/python-docx/issues/74

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    # Create a w:rStyle element, note this currently does not add the hyperlink style as its not in
    # the default template, I have left it here in case someone uses one that has the style in it
    rStyle = docx.oxml.shared.OxmlElement('w:rStyle')
    rStyle.set(docx.oxml.shared.qn('w:val'), 'Hyperlink')

    # Join all the xml elements together add add the required text to the w:r element
    rPr.append(rStyle)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink