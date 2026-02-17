from pathlib import Path
from contextlib import contextmanager
import re
import requests
import yaml

# noinspection PyPackageRequirements
from docx import Document
# noinspection PyPackageRequirements
from docx.shared import Cm
# noinspection PyPackageRequirements
from docx.oxml.ns import qn
# noinspection PyPackageRequirements
from docx.oxml import OxmlElement

from deplane.md_to_docx import insert_markdown, format_table


# mark translations; actual translation later when we know language"""
_ = lambda x: x


OBLIGATION = {  # convert required to Obligation
    'mandatory': _('Mandatory'),
    'conditional': _('Mandatory, conditional'),
    'optional': _('Optional'),
}

OCCURENCE = [
    _('Single'),
    _('Repeatable'),
]

VALIDATION = [
    _('This field must not be empty'),
]

FORMAT_TYPE = {  # convert datastore_type to Format Type
    'bigint': _('Integer'),
    'int': _('Integer'),
    'year': _('Integer'),
    'month': _('Integer'),
    'numeric': _('Numeric'),
    'money': _('Numeric'),
    'text': _('Text'),
    '_text': _('Text Array'),
    'text array': _('Text Array'),
    'date': _('Date'),
    'timestamp': _('Timestamp'),
}

STRINGS = [
    _('Describes the condition or conditions according to which a value shall be present in English. Indicates what the system will accept in this field.'),
    _('Describes the condition or conditions according to which a value shall be present in French. Indicates what the system will accept in this field.'),
]

NAP_MATCH = re.compile(r'^nap(\d)*$')


def write_docx(schema, filename, trans, lang, french_trans, remote_choices = None):
    """
    :param schema: recombinant-schema json dict
    :param filename: output docx filename to create
    :param trans: gettext translation object for language selected
    :param lang: language code selected (en/fr)
    """
    _ = trans.gettext

    document = Document(Path(__file__).parent / 'ENG_2_Colour.docx')
    section = document.sections[0]

    title = schema['title']
    if isinstance(title, dict):
        title = title[lang]

    section.different_first_page_header_footer = True
    section.first_page_header.paragraphs[0].text = ''
    section.header.paragraphs[0].text = _('Data Element Profile') + ' : ' + title

    # replace existing "TITLE" / "Subtitle" text
    document.paragraphs[0].runs[0].font.size = Cm(1.2)  # template title too large
    document.paragraphs[0].runs[0].text = _('Data Element Profile')
    document.paragraphs[0].runs[2].text = title
    # clear out the rest of the example text
    for para in reversed(document.paragraphs[2:]):
        delete_paragraph(para)

    #document.add_heading(_('Overview'), 1)
    #document.sections[-1].right_margin = Cm(2.54)  # restore right margin
    #insert_markdown(document, _(
    #    'The purpose of this document is to provide supplemental information that is '
    #    'not provided in the *Centralized Contract Publishing System: Training Guide*. '
    #    'It will provide information to users on data elements within the Quarterly '
    #    'Contracts template.'
    #))

    if schema.get('front_matter'):
        insert_markdown(document, schema['front_matter'][lang])

    document.add_heading(_('Legend'), 1)
    document.add_paragraph(_(
        'The following sample table provides a description of each field you will see '
        'for all elements:'
    ))

    with build_table(
            document,
            [Cm(4.69), Cm(11.80)],
            top_color='d9d9d9',
            left_color='c6d9f1') as (trow, mrow, ttable):

        trow(_('Attribute'), _('Attribute Description'))
        trow(
            _('Field Name EN'),
            _('This text should correspond directly with the field name in your template in English'),
        )
        trow(
            _('Field Name FR'),
            _('This text should correspond directly with the field name in your template in French'),
        )
        trow(
            _('Description EN'),
            _('This provides a brief description of the element in English'),
        )
        trow(
            _('Description FR'),
            _('This provides a brief description of the element in French'),
        )
        mrow(
            _('Obligation'),
            _('''
Indicates whether the element is required to always or sometimes be present 
(i.e., contain a value). Options are:

- Mandatory
- Mandatory, conditional
- Optional'''),
        )
        mrow(
            _('Format Type'),
            _('''
Options are:

- Integer (e.g. page count, year or month number)
- Numeric (e.g. decimal, currency values)
- Text
- Text Array (e.g. one or more codes from a controlled list)
- Date (YYYY-MM-DD)
- Timestamp (YYYY-MM-DD hh:mm:ss)''')
        )
        trow(
            _('Validation EN'),
            _(
'''Describes the condition or conditions according to which a value shall be present in English.
Indicates what the system will accept in this field.'''),
        )
        trow(
            _('Validation FR'),
            _(
'''Describes the condition or conditions according to which a value shall be present in French.
Indicates what the system will accept in this field.'''),
        )
        trow(
            _('Example Value'),
            _(
                'Provide one or more real examples of the values that may appear, '
                'e.g. “CODE1” or “Family Services Reform Program”'
            ),
        )

    document.add_paragraph(_('\nControlled List Values:'))
    with build_table(
            document,
            [Cm(3.69), Cm(6.4), Cm(6.4)],
            top_color='d9d9d9') as (trow, mrow, ttable):
        trow(_('Code'), _('English'), _('French'))
        trow('CODE1', _('English Description 1'), _('French Description 1'))
        trow('CODE2', _('English Description 2'), _('French Description 2'))

    document.add_page_break()

    for rnum, res in enumerate(schema['resources'], 1):
        res_title = schema['title']
        if isinstance(res_title, dict):
            res_title = res_title[lang]

        document.add_heading(res_title + '\n', 1)

        rname = res.get('resource_name', '')
        is_national_action_plan = NAP_MATCH.match(rname) is not None
        nap_version = None
        if is_national_action_plan:
            nap_version = int(rname.replace('nap', ''))

        for fnum, field in enumerate(res['fields'], 1):
            typ = field.get('type', field.get('datastore_type'))
            fid = field.get('id', field.get('datastore_id'))
            is_pub_field = field.get('published_resource_computed_field', False)
            if not field.get('import_template_include', True):
                continue
            if not field.get('visible_to_public', True):
                continue
            if fid == 'owner_org' or fid == 'owner_org_title':
                continue

            flabel = field['label'] if isinstance(field['label'], str) else field['label'][lang]
            document.add_heading(f'{rnum}-{fnum} {flabel}\n', 2)

            with build_table(
                    document,
                    [Cm(4.69), Cm(11.80)],
                    top_color='d9d9d9',
                    left_color='c6d9f1') as (trow, mrow, ttable):
                
                trow(_('Attribute'), _('Attribute Description'))

                flabel_en = field['label'] if isinstance(field['label'], str) else field['label']['en']
                flabel_fr = field['label'] if isinstance(field['label'], str) else field['label']['fr']
                if flabel_fr == flabel_en:
                    flabel_fr = ''
                trow(_('Field Name EN'), flabel_en)
                trow(_('Field Name FR'), flabel_fr)

                trow(_('ID'), fid)

                mrow(_('Description EN'), field.get('description', {}).get('en', ''))
                mrow(_('Description FR'), field.get('description', {}).get('fr', ''))

                # JSON Schema vs YAML Schema obligations gets a bit complicated...
                obligation_en = field.get('obligation', 'optional')
                if not isinstance(obligation_en, dict) and obligation_en in OBLIGATION:
                    obligation_en = trans.gettext(OBLIGATION[obligation_en])
                elif isinstance(obligation_en, dict):
                    obligation_en = obligation_en.get('en')
                    if obligation_en in OBLIGATION:
                        obligation_en = trans.gettext(OBLIGATION[obligation_en])
                    else:
                        obligation_en = trans.gettext(obligation_en)
                if not obligation_en:
                    obligation_en = trans.gettext('Optional')
                obligation_fr = field.get('obligation', 'optional')
                if not isinstance(obligation_fr, dict) and obligation_fr in OBLIGATION:
                    obligation_fr = french_trans.gettext(OBLIGATION[obligation_fr])
                elif isinstance(obligation_fr, dict):
                    obligation_fr = obligation_fr.get('fr')
                    if obligation_fr in OBLIGATION:
                        obligation_fr = french_trans.gettext(OBLIGATION[obligation_fr])
                    else:
                        obligation_fr = french_trans.gettext(obligation_fr)
                elif isinstance(obligation_fr, str) and obligation_fr in OBLIGATION.values():
                    obligation_fr = french_trans.gettext(obligation_fr)
                if not obligation_fr:
                    obligation_fr = french_trans.gettext('Optional')
                mrow(_('Obligation EN'), obligation_en)
                mrow(_('Obligation FR'), obligation_fr)

                trow(_('Occurrence'), trans.gettext(field.get('occurrence', 'Single')))

                mrow(_('Format Type'), trans.gettext(FORMAT_TYPE[typ]))

                # There are some basic string validations instead of en/fr dicts
                validation_en = field.get('validation', '')
                if isinstance(validation_en, dict):
                    validation_en = validation_en.get('en', '')
                elif isinstance(validation_en, str) and validation_en in VALIDATION:
                    validation_en = trans.gettext(validation_en)
                elif isinstance(validation_en, str):
                    validation_en = validation_en
                validation_fr = field.get('validation', '')
                if isinstance(validation_fr, dict):
                    validation_fr = validation_fr.get('fr', '')
                elif isinstance(validation_fr, str) and validation_fr in VALIDATION:
                    validation_fr = french_trans.gettext(validation_fr)
                elif isinstance(validation_fr, str):
                    validation_fr = validation_fr
                trow(_('Validation EN'), validation_en)
                trow(_('Validation FR'), validation_fr)

                if field.get('character_limit', field.get('max_chars')):
                    trow(_('Character Limit'), str(field.get('character_limit', field.get('max_chars'))))

                eg_r = res.get('example_record' ,res.get('examples', {}).get('record', None))
                if not is_pub_field and eg_r.get(fid):
                    eg = eg_r[fid]
                    if isinstance(eg, list):
                        eg = ','.join(eg)
                    trow(_('Example Value'), str(eg))
                    
            fchoices = field.get('choices')
            if 'choices_file' in field:
                cf = remote_choices + '/' + field['choices_file']
                rf = requests.get(cf)
                if cf.endswith('.json'):
                    fchoices = rf.json()
                elif cf.endswith('.yaml') or cf.endswith('.yml'):
                    fchoices = yaml.safe_load(rf.content)
                print('Fetching choices for %s from Github remote: %s' % (fid, cf))
            if fchoices:
                document.add_paragraph(_('\nControlled List Values:'))
                with build_table(
                        document,
                        [Cm(3.69), Cm(6.4), Cm(6.4)],
                        top_color='d9d9d9') as (trow, mrow, ttable):
                    trow(_('Code'), _('English'), _('French'))
                    for c, v in fchoices.items():
                        if isinstance(v, dict):
                            trow(c, v.get('en', ''), v.get('fr', ''))
                        else:
                            trow(c, v, v)
                        # Special stuff for National Action Plans
                        if (
                          is_national_action_plan and 
                          isinstance(v, dict) and
                          (
                           'due_date' in v or
                           'deadline' in v or
                           'lead_dept' in v or
                           's4d' in v  
                          )
                        ):
                            with ttable([Cm(3.69), Cm(12.8)], top_color='ddd8e8', left_color='e3e3e3') as innerT:
                                cells = innerT.add_row().cells
                                cells[0].text = _('Extras for: %s') % c
                                cells[1].text = _('Extra Value')

                                if nap_version == 5:
                                    cells = innerT.add_row().cells
                                    cells[0].text = _('Due Date')
                                    default = 'TODO - supply value (must be one of reporting_period values)'
                                    cells[1].text = v.get('due_date') or default

                                cells = innerT.add_row().cells
                                cells[0].text = _('Deadline (English)')
                                default = 'TODO - supply value (semantics e.g. "First of June 2026")'
                                cells[1].text = v.get('deadline', {}).get('en') or default

                                cells = innerT.add_row().cells
                                cells[0].text = _('Deadline (French)')
                                default = 'TODO - supply value (semantics e.g. "First of June 2026")'
                                cells[1].text = v.get('deadline', {}).get('fr') or default

                                cells = innerT.add_row().cells
                                cells[0].text = _('Lead Department')
                                default = 'TODO - supply value (must be organization abbreviation found here: https://open.canada.ca/data/en/api/action/organization_list)'
                                lead = v.get('lead_dept')
                                if not lead:
                                    lead = default
                                elif isinstance(lead, list):
                                    lead = ','.join(lead)
                                cells[1].text = lead

                                if nap_version == 5:
                                    cells = innerT.add_row().cells
                                    cells[0].text = _('Summit for Democracy')
                                    default = 'TODO - supply value (must be "true" or "false")'
                                    cells[1].text = v.get('s4d') or default
                        # :END: Special stuff for National Action Plans :END:

            document.add_paragraph('\n\n')

        document.add_page_break()

    document.save(filename)


def delete_paragraph(para):
    p = para._element
    p.getparent().remove(p)
    p._p = p._element = None


@contextmanager
def build_table(document, widths, **format_args):
    """
    Context manager for building a table with text rows and markdown rows

    Yields (trow, mrow) functions for adding text/markdown rows to the table
    """
    table = document.add_table(rows=0, cols=len(widths))
    table.autofit = False
    def trow(first, *rest):
        """Add a text row to table"""
        cells = table.add_row().cells
        cells[0].text = first
        for i, c in enumerate(rest, 1):
            cells[i].text = c

    def mrow(first, *rest):
        """Add a markdown row to table"""
        cells = table.add_row().cells
        cells[0].text = first
        for i, c in enumerate(rest, 1):
            if c:
                delete_paragraph(cells[i].paragraphs[0])
            insert_markdown(cells[i], c)

    @contextmanager
    def ttable(inner_widths, **inner_format_args):
        """Add a nested table"""
        cells = table.add_row().cells
        cell_1 = cells[0]
        cell_n = cells[-1]
        merged_cell = cell_1.merge(cell_n)
        merged_cell._element.clear_content()  # clear leading paragraph
        innerTable = merged_cell.add_table(rows=0, cols=len(inner_widths))
        set_cell_margins(merged_cell,
                         top=Cm(0),
                         bottom=Cm(0),
                         start=Cm(0),
                         end=Cm(0))
        yield innerTable
        format_table(innerTable, widths=inner_widths, **inner_format_args)

    yield trow, mrow, ttable

    format_table(table, widths=widths, **format_args)


def set_cell_margins(cell, **kwargs):
    """
    Set the margins for a table cell.
    
    Values are specified in twentieths of a point (1/1440 of a centimeter).
    Use the docx.shared.Cm for convenience.

    Usage:
    set_cell_margins(cell, top=Cm(0), bottom=Cm(0), start=Cm(0), end=Cm(0))
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m in ["top", "start", "bottom", "end"]:
        if m in kwargs:
            node = OxmlElement(f"w:{m}")
            # Ensure the value is an integer representing twentieths of a point
            value_dxa = int(kwargs[m].emu / 6350)
            node.set(qn('w:w'), str(value_dxa))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
    tcPr.append(tcMar)
